"""Service for parsing resumes and extracting text from files."""
import os
import subprocess
import shutil
import tempfile
from io import BytesIO
from typing import Optional
from docx import Document
import PyPDF2

from app.utils.logging import get_logger
from app.utils.safe_logger import safe_extra
from app.utils.cleaning import normalize_text

logger = get_logger(__name__)

# Try to import Apache Tika for .doc file support (PRIMARY METHOD - Currently Working)
try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False
    logger.warning("Apache Tika not available. .doc files cannot be processed without Tika.")

# Try to import olefile for .doc file support (fallback method)
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    logger.debug("olefile not available. Will use Tika only for .doc files.")

# Check for antiword command-line tool
ANTIWORD_AVAILABLE = shutil.which("antiword") is not None
if not ANTIWORD_AVAILABLE:
    logger.debug("antiword command not found. Install it for better .doc extraction.")

# Check for LibreOffice command-line tool
LIBREOFFICE_AVAILABLE = shutil.which("soffice") is not None or shutil.which("libreoffice") is not None
if not LIBREOFFICE_AVAILABLE:
    logger.debug("LibreOffice not found. Install it for most reliable .doc conversion.")


class ResumeParser:
    """Service for parsing resume files and extracting text."""
    
    def __init__(self):
        """Initialize ResumeParser."""
        pass
    
    async def extract_text(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from uploaded file based on extension.
        
        Args:
            file_content: The binary content of the file
            filename: Name of the file (used to determine file type)
        
        Returns:
            Extracted text content as string
        
        Raises:
            ValueError: If file type is not supported or extraction fails
        """
        try:
            if filename.lower().endswith('.pdf'):
                return self._extract_pdf_text(file_content)
            elif filename.lower().endswith('.docx'):
                return self._extract_docx_text(file_content)
            elif filename.lower().endswith('.doc'):
                return self._extract_doc_text(file_content)
            elif filename.lower().endswith('.txt'):
                return file_content.decode('utf-8', errors='ignore')
            else:
                # Try as text for unknown extensions
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            # Safe logging - avoid using reserved LogRecord attributes
            error_msg = f"Error extracting text from {filename}: {e}"
            print(f"[ERROR] {error_msg}")  # Print to console for visibility
            try:
                # Use safe_extra to prevent LogRecord conflicts
                safe_extras = safe_extra({"error": str(e), "file_name": filename})
                logger.error(error_msg, extra=safe_extras)
            except Exception as log_error:
                # If logging fails, at least print the error
                print(f"[CRITICAL] Logging failed: {log_error}")
                print(f"[CRITICAL] Original error: {error_msg}")
            raise ValueError(f"Failed to extract text from file: {e}")
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            raw_text = "\n".join(text_parts)
            # Normalize whitespace (remove extra spaces, normalize line breaks)
            normalized_text = normalize_text(raw_text) or raw_text
            return normalized_text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}", extra={"error": str(e)})
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file, including headers, footers, and tables.
        Uses python-docx first, then falls back to Apache Tika for comprehensive extraction.
        """
        # First try: Use python-docx (fast and reliable for most cases)
        try:
            doc_file = BytesIO(file_content)
            doc = Document(doc_file)
            text_parts = []
            
            # Extract from main document paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables (contact info is sometimes in tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract from headers (contact info is often in headers)
            for section in doc.sections:
                # Header text
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    # Headers can also have tables
                    for table in section.header.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                
                # Footer text
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    # Footers can also have tables
                    for table in section.footer.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_text.append(cell_text)
                            if row_text:
                                text_parts.append(" | ".join(row_text))
            
            raw_text = "\n".join(text_parts)
            # Normalize whitespace (remove extra spaces, normalize line breaks)
            normalized_text = normalize_text(raw_text) or raw_text
            
            # If we found text, return it
            if normalized_text.strip():
                return normalized_text
        except Exception as e:
            logger.debug(f"python-docx extraction had issues: {e}")
        
        # Second try: Use Apache Tika for comprehensive extraction (includes text boxes, headers, footers)
        # Tika can extract more content including text boxes and complex layouts
        if TIKA_AVAILABLE:
            try:
                # Create temporary file for Tika
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(file_content)
                    temp_docx_path = temp_file.name
                
                try:
                    logger.debug("Attempting DOCX extraction using Apache Tika for comprehensive extraction")
                    parsed = tika_parser.from_file(temp_docx_path)
                    if parsed and 'content' in parsed and parsed['content']:
                        text = parsed['content'].strip()
                        if text:
                            # Normalize whitespace
                            normalized_text = normalize_text(text) or text
                            logger.info(f"Successfully extracted DOCX using Apache Tika: {len(normalized_text)} characters")
                            return normalized_text
                finally:
                    # Clean up temp file
                    if os.path.exists(temp_docx_path):
                        os.unlink(temp_docx_path)
            except Exception as tika_error:
                logger.debug(f"Apache Tika extraction failed: {tika_error}")
        
        # If both methods failed or returned empty, raise error
        raise ValueError("Failed to extract text from DOCX file")
    
    def _extract_doc_text(self, file_content: bytes) -> str:
        """
        Extract text from DOC file (older Microsoft Word format).
        Uses methods in order of reliability:
        1. Apache Tika (PRIMARY - Currently Working)
        2. LibreOffice headless conversion (if available)
        3. antiword (if available)
        4. python-docx fallback (might work for some files)
        5. olefile (basic binary extraction - fallback)
        """
        # Create temporary file for .doc content
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
            temp_file.write(file_content)
            temp_doc_path = temp_file.name
        
        try:
            # Method 1: Apache Tika (PRIMARY METHOD - Currently Working)
            if TIKA_AVAILABLE:
                try:
                    tika_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$  USING APACHE TIKA  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Extracting text from .doc file using Apache Tika\n"
                        "$$$  This is the PRIMARY method for processing .doc files\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(tika_msg)
                    logger.info(tika_msg)
                    parsed = tika_parser.from_file(temp_doc_path)
                    if parsed and 'content' in parsed and parsed['content']:
                        text = parsed['content'].strip()
                        if text:
                            # Normalize whitespace (remove extra spaces, normalize line breaks)
                            normalized_text = normalize_text(text) or text
                            success_msg = (
                                "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                "$$$$$$$$$$$$$$$$$$$$$$  APACHE TIKA SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                f"$$$  Successfully extracted {len(normalized_text)} characters using Apache Tika\n"
                                "$$$  METHOD USED: Apache Tika (tika-python library)\n"
                                "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                            )
                            print(success_msg)
                            logger.info(
                                success_msg,
                                extra={"extraction_method": "apache_tika", "text_length": len(normalized_text)}
                            )
                            return normalized_text
                except Exception as tika_error:
                    error_msg = f"Apache Tika extraction failed: {tika_error}"
                    print(f"[WARNING] {error_msg}")
                    logger.warning(error_msg)
            
            # Method 2: LibreOffice headless conversion (if available)
            if LIBREOFFICE_AVAILABLE:
                try:
                    lo_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$  USING LIBREOFFICE  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Converting .doc to .docx using LibreOffice headless\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(lo_msg)
                    logger.info(lo_msg)
                    # Create temp directory for output
                    with tempfile.TemporaryDirectory() as temp_dir:
                        # Use LibreOffice to convert .doc to .docx
                        libreoffice_cmd = shutil.which("soffice") or shutil.which("libreoffice")
                        cmd = [
                            libreoffice_cmd,
                            "--headless",
                            "--convert-to", "docx",
                            "--outdir", temp_dir,
                            temp_doc_path
                        ]
                        
                        result = subprocess.run(
                            cmd,
                            capture_output=True,
                            text=True,
                            timeout=30
                        )
                        
                        if result.returncode == 0:
                            # Find the converted .docx file
                            converted_docx = os.path.join(temp_dir, os.path.basename(temp_doc_path).replace('.doc', '.docx'))
                            if os.path.exists(converted_docx):
                                # Read the converted .docx and extract text
                                with open(converted_docx, 'rb') as f:
                                    docx_content = f.read()
                                text = self._extract_docx_text(docx_content)
                                if text.strip():
                                    success_msg = (
                                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                        "$$$$$$$$$$$$$$$$$$$$$$  LIBREOFFICE SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                                        f"$$$  Successfully extracted {len(text)} characters using LibreOffice\n"
                                        "$$$  METHOD USED: LibreOffice (converted .doc to .docx, then extracted)\n"
                                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                                    )
                                    print(success_msg)
                                    logger.info(
                                        success_msg,
                                        extra={"extraction_method": "libreoffice", "text_length": len(text)}
                                    )
                                    return text
                except subprocess.TimeoutExpired:
                    logger.warning("LibreOffice conversion timed out")
                except Exception as lo_error:
                    logger.debug(f"LibreOffice conversion failed: {lo_error}")
            
            # Method 3: antiword (if available)
            if ANTIWORD_AVAILABLE:
                try:
                    antiword_msg = (
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$  USING ANTIWORD  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                        "$$$  Extracting text from .doc file using antiword\n"
                        "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                    )
                    print(antiword_msg)
                    logger.info(antiword_msg)
                    result = subprocess.run(
                        ["antiword", temp_doc_path],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        # Normalize whitespace (remove extra spaces, normalize line breaks)
                        normalized_text = normalize_text(result.stdout) or result.stdout
                        success_msg = (
                            "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                            "$$$$$$$$$$$$$$$$$$$$$$  ANTIWORD SUCCESS  $$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$\n"
                            f"$$$  Successfully extracted {len(normalized_text)} characters using antiword\n"
                            "$$$  METHOD USED: antiword (command-line tool)\n"
                            "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
                        )
                        print(success_msg)
                        logger.info(
                            success_msg,
                            extra={"extraction_method": "antiword", "text_length": len(normalized_text)}
                        )
                        return normalized_text
                except subprocess.TimeoutExpired:
                    logger.warning("antiword extraction timed out")
                except Exception as aw_error:
                    logger.debug(f"antiword extraction failed: {aw_error}")
            
            # Method 4: Try python-docx as fallback (might work for some .doc files that are actually .docx)
            try:
                logger.debug("Attempting .doc extraction using python-docx fallback")
                doc_file = BytesIO(file_content)
                doc = Document(doc_file)
                text_parts = []
                for paragraph in doc.paragraphs:
                    text_parts.append(paragraph.text)
                # Also try to extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                extracted_text = "\n".join(text_parts)
                if extracted_text.strip():
                    # Normalize whitespace (remove extra spaces, normalize line breaks)
                    normalized_text = normalize_text(extracted_text) or extracted_text
                    logger.info("Successfully extracted .doc file using python-docx fallback")
                    return normalized_text
            except Exception as fallback_error:
                logger.debug(f"python-docx fallback failed: {fallback_error}")
            
            # Method 5: Try to extract using olefile (for binary .doc files - last resort)
            if OLEFILE_AVAILABLE:
                try:
                    logger.debug("Attempting .doc extraction using olefile (last resort)")
                    # .doc files are OLE compound documents
                    ole = olefile.OleFileIO(BytesIO(file_content))
                    # Try to find WordDocument stream
                    if ole.exists('WordDocument'):
                        stream = ole.openstream('WordDocument')
                        # Read and try to extract text (basic extraction)
                        data = stream.read()
                        # Simple text extraction from binary data
                        # This is a basic approach - look for readable text
                        text_chunks = []
                        current_chunk = b""
                        for byte in data:
                            if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                                current_chunk += bytes([byte])
                            else:
                                if len(current_chunk) > 3:
                                    try:
                                        text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                                    except:
                                        pass
                                current_chunk = b""
                        if current_chunk:
                            try:
                                text_chunks.append(current_chunk.decode('ascii', errors='ignore'))
                            except:
                                pass
                        ole.close()
                        extracted_text = "\n".join(text_chunks)
                        if extracted_text.strip():
                            # Normalize whitespace (remove extra spaces, normalize line breaks)
                            normalized_text = normalize_text(extracted_text) or extracted_text
                            logger.info("Successfully extracted .doc file using olefile")
                            return normalized_text
                    ole.close()
                except Exception as ole_error:
                    logger.debug(f"olefile extraction failed: {ole_error}")
            
            # If all methods fail, raise an error with helpful message
            raise ValueError(
                "Failed to extract text from .doc file using all available methods.\n"
                "Installation options:\n"
                "1. Apache Tika: pip install tika (REQUIRES Java runtime) - PRIMARY METHOD\n"
                "2. LibreOffice (headless): Most reliable for production\n"
                "   - Windows: Download from https://www.libreoffice.org/\n"
                "   - Linux: sudo apt-get install libreoffice\n"
                "3. antiword: Good for plain-text extraction\n"
                "   - Windows: Download from http://www.winfield.demon.nl/\n"
                "   - Linux: sudo apt-get install antiword\n"
                "4. olefile: pip install olefile (basic support, already installed)\n"
                "5. Convert .doc files to .docx format before processing"
            )
        finally:
            # Clean up temp file
            if os.path.exists(temp_doc_path):
                os.unlink(temp_doc_path)

